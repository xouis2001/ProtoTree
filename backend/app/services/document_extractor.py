from io import BytesIO

from docx import Document
from pypdf import PdfReader


SUPPORTED_PROTOCOL_EXTENSIONS = {".pdf", ".docx"}


def normalize_extracted_text(text: str) -> str:
    paragraphs = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    blocks: list[str] = []
    pending: list[str] = []
    for paragraph in paragraphs:
        if paragraph:
            pending.append(paragraph)
        elif pending:
            blocks.append(" ".join(pending))
            pending = []
    if pending:
        blocks.append(" ".join(pending))
    return "\n\n".join(blocks).strip()


def extract_pdf_text(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    return normalize_extracted_text("\n\n".join(pages))


def extract_docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_rows: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_rows.append(" | ".join(cells))
    return normalize_extracted_text("\n".join([*paragraphs, *table_rows]))


def extract_document_text(filename: str, content: bytes) -> str:
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if suffix == "pdf":
        return extract_pdf_text(content)
    if suffix == "docx":
        return extract_docx_text(content)
    raise ValueError("Only PDF and DOCX files are supported")


def build_plain_protocol_structured(title: str, text: str) -> dict:
    return {
        "experiment_name": title,
        "experiment_type": "",
        "experiment_subtype": "",
        "content": text,
        "steps": [],
    }
